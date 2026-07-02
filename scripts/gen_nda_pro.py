"""NDA Generator — Unified Taiwan Contract Formatting Generator

Usage:
    python scripts/gen_nda_pro.py                        # Default: ~/Desktop/NDA.docx
    python scripts/gen_nda_pro.py --output ./my-nda.docx  # Custom path
    python scripts/gen_nda_pro.py --compact               # Compact mode
    python scripts/gen_nda_pro.py --party-a "XXX Corp"    # Custom party name

Consolidates: gen_nda_compact.py, gen_nda_docx.py, gen_nda_docx_standalone.py,
             gen_nda_final.py, gen_nda_v2.py (now deprecated → import this)

Formatting specs (verified against real Taiwan contract practice):
- Font: 標楷體 12pt (TW-Kai)
- Line spacing: Fixed 22pt (compact=18pt)
- Alignment: Justified (左右對齊)
- Margins: 2.5cm all sides (compact=2.0cm)
- Party block (甲方/乙方): Borderless table for perfect colon alignment
- Article hierarchy: 第一條, 一、, (一), 1.
- Signature: Separate page, dual-column borderless table
- Page number: Bottom center
"""
import argparse
import os
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def build_nda(output_path: str, compact: bool = False,
              party_a: str = "甲方", party_b: str = "乙方") -> str:
    """Generate NDA .docx with Taiwan contract formatting.

    Returns the absolute output path.
    """
    FONT = "TW-Kai"
    FS = 12 if not compact else 11
    LINE_H = Pt(22) if not compact else Pt(18)
    MARGIN = 2.5 if not compact else 2.0

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for a in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']:
        setattr(sec, a, Cm(MARGIN))

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(FS)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = style.paragraph_format
    pf.line_spacing = LINE_H
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def borderless_table(rows, cols):
        t = doc.add_table(rows=rows, cols=cols)
        t.style = 'Table Grid'
        for row in t.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                b = OxmlElement('w:tcBorders')
                for bn in ['top', 'left', 'bottom', 'right']:
                    e = OxmlElement(f'w:{bn}')
                    e.set(qn('w:val'), 'none')
                    e.set(qn('w:sz'), '0')
                    e.set(qn('w:space'), '0')
                    e.set(qn('w:color'), 'auto')
                    b.append(e)
                tcPr.append(b)
        return t

    def add(text, bold=False, size=FS, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_b=0, space_a=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_b)
        p.paragraph_format.space_after = Pt(space_a)
        p.paragraph_format.line_spacing = LINE_H
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(size)
        r.bold = bold
        r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        return p

    def sec_head(text):
        return add(text, bold=True, size=FS, space_b=6)

    # === TITLE ===
    add('雙向保密協議', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add('Non-Disclosure Agreement', size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_b=8)

    # === PARTY BLOCK ===
    t = borderless_table(3, 2)
    for r in range(3):
        for c in range(2):
            cell = t.cell(r, c)
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tbl_data = [
        ('立協議人：', True, '', ''),
        (f'{party_a}：', False, '（以下簡稱甲方）', False),
        (f'{party_b}：', False, '（以下簡稱乙方）', False),
    ]
    for i, (txt, bld, txt2, bld2) in enumerate(tbl_data):
        p0 = t.cell(i, 0).paragraphs[0]
        r0 = p0.add_run(txt)
        r0.font.name = FONT
        r0.font.size = Pt(FS)
        r0.bold = bld
        r0.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        if txt2:
            p1 = t.cell(i, 1).paragraphs[0]
            r1 = p1.add_run(txt2)
            r1.font.name = FONT
            r1.font.size = Pt(FS)
            r1.bold = bld2
            r1.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    add('緣甲方與乙方為評估及建立商業合作關係（以下簡稱「目的」），'
        '就雙方於合作期間所交換之保密資訊，爰協議如下：', space_b=4)

    # === ARTICLES ===
    articles = [
        ('第一條　定義', [
            '一、保密資訊：指甲方或乙方（以下簡稱「揭露方」）於目的範圍內，'
            '以書面、電子、圖像或其他可留存形式，向他方（以下簡稱「接收方」）'
            '揭露並標示為機密或保密之資訊。',
            '二、目的：指雙方因合作、洽談及執行專案所進行之一切評估、討論及業務往來。',
        ]),
        ('第二條　保密義務', [
            '一、接收方應以合理之注意程度，保護揭露方之保密資訊，'
            '並僅得為目的範圍內使用該保密資訊。',
            '二、接收方得使其因執行目的而需知悉之董事、監察人、經理人、員工'
            '（以下合稱「必要人員」）接觸保密資訊，'
            '但應確保該等人員知悉並遵守本協議之保密義務。',
            '三、接收方對其必要人員違反本協議之行為，應負連帶責任。',
        ]),
        ('第三條　例外條款', [
            '下列情形不構成保密資訊，接收方不負保密義務：',
            '（一）接收方於揭露時能證明已明知之資訊；',
            '（二）非因接收方違反本協議而成為公開資訊；',
            '（三）接收方自無保密義務之第三人合法取得之資訊；',
            '（四）揭露方事前書面同意揭露或公開之資訊；',
            '（五）依法律、法院命令、政府機關要求而須揭露之資訊，'
            '接收方應於可行範圍內事先通知揭露方，'
            '並僅揭露依法令要求之最小範圍。',
        ]),
        ('第四條　資訊返還與銷毀', [
            '一、任一方於目的完成後或依他方書面要求時，'
            '應於三十日內返還或銷毀其所持有之保密資訊（含所有副本），'
            '並出具書面證明。',
            '二、前項規定不適用於接收方依其內部備份政策或法規要求而留存之備份檔案，'
            '惟該留存檔案仍應繼續受本協議之保密義務拘束。',
        ]),
        ('第五條　生效與期間', [
            '一、本協議自最後簽署日起生效（以下簡稱「協議生效日」）。',
            '二、任一方得隨時以三十日書面預告通知他方終止本協議，'
            '惟終止前已揭露之保密資訊不受影響。',
            '三、各筆保密資訊之保密義務，自該資訊首次揭露日起算三年。',
            '四、協議終止後，第四條（返還與銷毀）與第六條（違約責任）繼續有效。',
        ]),
        ('第六條　違約責任', [
            '一、任一方違反本協議之保密義務，應賠償他方因此所受之損害。',
            '二、雙方同意，任一方如違反第二條或第三條之規定，'
            '應按次給付違約金新臺幣五十萬元；若違約情節連續或持續發生，'
            '以日計，每日另計新臺幣五萬元。',
            '三、若實際損害逾前項違約金之金額，他方仍得請求超出部分。',
        ]),
        ('第七條　管轄與準據法', [
            '一、本協議之解釋、效力及履行，以中華民國法律為準據法。',
            '二、因本協議所生之爭議，雙方合意以臺灣臺北地方法院為第一審管轄法院。',
        ]),
        ('第八條　其他條款', [
            '一、本協議構成雙方就保密事項之完整合意，取代先前行為或口頭約定。',
            '二、本協議之修改應經雙方書面同意。',
            '三、本協議任一條款若被認定為無效或無法執行，不影響其他條款之效力。',
        ]),
    ]

    for title, lines in articles:
        sec_head(title)
        for line in lines:
            add(line)

    # === SIGNATURE (separate page) ===
    doc.add_page_break()
    add('立協議人', bold=True, size=FS, align=WD_ALIGN_PARAGRAPH.CENTER)
    add('')

    sig = borderless_table(1, 2)
    for side_idx, (party, fields) in enumerate([
        (party_a, [f'{party_a}：（簽章）', '代表人：', '統一編號：', '地址：', '電話：']),
        (party_b, [f'{party_b}：（簽章）', '代表人：', '統一編號：', '地址：', '電話：']),
    ]):
        cell = sig.cell(0, side_idx)
        cell.text = ''
        for i, f in enumerate(fields):
            pp = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            pp.paragraph_format.space_before = Pt(4)
            pp.paragraph_format.space_after = Pt(0)
            rr = pp.add_run(f)
            rr.font.name = FONT
            rr.font.size = Pt(FS)
            rr.bold = (i == 0)
            rr.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    add('')
    add('中　華　民　國　　　年　　　月　　　日',
        align=WD_ALIGN_PARAGRAPH.CENTER)

    # Footer page number
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    rr = fp.add_run()
    rr.font.size = Pt(9)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' PAGE ')
    run_e = OxmlElement('w:r')
    t_e = OxmlElement('w:t')
    t_e.text = '1'
    run_e.append(t_e)
    fld.append(run_e)
    fp._p.append(fld)

    abs_out = os.path.abspath(output_path)
    doc.save(abs_out)
    return abs_out


def main():
    default_out = os.path.join(os.path.expanduser("~"), "Desktop", "NDA.docx")
    parser = argparse.ArgumentParser(
        description="智研法律 — NDA 合約產生器（台灣法務排版）")
    parser.add_argument("--output", "-o", default=default_out,
                        help=f"輸出路徑（預設：{default_out}）")
    parser.add_argument("--compact", "-c", action="store_true",
                        help="緊湊模式（11pt/18pt行距/2.0cm邊距）")
    parser.add_argument("--party-a", default="甲方",
                        help="甲方名稱（預設：甲方）")
    parser.add_argument("--party-b", default="乙方",
                        help="乙方名稱（預設：乙方）")
    args = parser.parse_args()

    out = build_nda(args.output, compact=args.compact,
                    party_a=args.party_a, party_b=args.party_b)
    print(f'✅ {out}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
