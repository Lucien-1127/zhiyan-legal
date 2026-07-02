"""Generate NDA v1.2 FINAL - standalone version (no zhiyan_legal dep)"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── Court-compliant params (民事訴訟書狀規則 §3) ──────────
A4_W = 21.0
A4_H = 29.7
MARGIN = 2.5
FONT = "TW-Kai"
FS = 14
LINE = Pt(28)

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(A4_W)
sec.page_height = Cm(A4_H)
sec.top_margin = Cm(MARGIN)
sec.bottom_margin = Cm(MARGIN)
sec.left_margin = Cm(MARGIN)
sec.right_margin = Cm(MARGIN)

style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(FS)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
style.paragraph_format.line_spacing = LINE

def add_para(text, bold=False, align=None, size=None, indent=None, spacing_before=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if spacing_before:
        p.paragraph_format.space_before = spacing_before
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size or FS)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    return p

def add_section_title(text):
    p = add_para(text, bold=True, size=16, spacing_before=Pt(12))
    return p

# ══════════ DOCUMENT BODY ══════════
add_para('雙向保密協議', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Non-Disclosure Agreement (NDA) v1.2 FINAL', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('')

add_section_title('當事人')
add_para('揭露方：【                        】（以下簡稱「甲方」）')
add_para('接收方：【                        】（以下簡稱「乙方」）')
add_para('')
add_para('甲方與乙方（以下合稱「雙方」，單稱「一方」）為評估及建立商業合作關係（以下簡稱「目的」），就雙方於合作期間所交換之保密資訊，爰協議如下：', indent=0.7)

add_section_title('第一條　定義')
add_para('1.1　保密資訊：指甲方或乙方（以下簡稱「揭露方」）於目的範圍內，以書面、電子、圖像或其他可留存形式，向他方（以下簡稱「接收方」）揭露並標示為機密或保密之資訊。以口頭或視覺方式揭露之資訊，不構成保密資訊，但揭露方得於揭露後十四日內以書面摘要確認，自書面送達接收方時起，該摘要內容即視為保密資訊，除非接收方於送達後七日內以書面提出具體異議。', indent=0.7)
add_para('1.2　目的：指雙方因合作、洽談及執行專案所進行之一切評估、討論及業務往來。', indent=0.7)

add_section_title('第二條　保密義務')
add_para('2.1　接收方應以合理之注意程度，保護揭露方之保密資訊，並僅得為目的範圍內使用該保密資訊。該注意程度不得低於保護自身同類型機密資訊之標準。', indent=0.7)
add_para('2.2　接收方得使其因執行目的而需知悉之董事、監察人、經理人、員工（以下合稱「必要人員」）接觸保密資訊，但應確保該等人員知悉並遵守本協議之保密義務。', indent=0.7)
add_para('2.3　接收方對其必要人員違反本協議之行為，應負連帶責任。', indent=0.7)

add_section_title('第三條　例外條款')
add_para('下列情形不構成保密資訊，接收方不負保密義務。接收方就主張例外之資訊，應負舉證責任：', indent=0.7)
add_para('3.1　接收方於揭露時能證明已明知之資訊；', indent=1.4)
add_para('3.2　非因接收方違反本協議而成為公開資訊；', indent=1.4)
add_para('3.3　接收方自無保密義務之第三人合法取得之資訊；', indent=1.4)
add_para('3.4　揭露方事前書面同意揭露或公開之資訊；', indent=1.4)
add_para('3.5　依法律、法院命令、政府機關要求而須揭露之資訊，惟接收方應於可行範圍內事先通知揭露方，並僅揭露依法令要求之最小範圍。', indent=1.4)

add_section_title('第四條　資訊返還與銷毀')
add_para('4.1　任一方於目的完成後或依他方書面要求時，應於三十日內返還或銷毀其所持有之保密資訊（含所有副本），並出具書面證明。', indent=0.7)
add_para('4.2　前項規定不適用於接收方依其內部備份政策或法規要求而留存之備份檔案，惟該留存檔案仍應繼續受本協議之保密義務拘束。', indent=0.7)

add_section_title('第五條　生效與期間')
add_para('5.1　本協議自最後簽署日起生效（以下簡稱「協議生效日」）。', indent=0.7)
add_para('5.2　任一方得隨時以三十日書面預告通知他方終止本協議（以下簡稱「協議終止日」），惟終止前已揭露之保密資訊不受影響。', indent=0.7)
add_para('5.3　各筆保密資訊之保密義務，自該資訊首次揭露日起算三年；但若該資訊構成營業秘密法所稱之營業秘密，則保密義務持續至該資訊依法不再構成營業秘密為止。', indent=0.7)
add_para('5.4　協議終止後，第四條（返還與銷毀）與第六條（違約責任）繼續有效。各筆資訊之保密義務期間依第5.3條定之。', indent=0.7)

add_section_title('第六條　違約責任')
add_para('6.1　任一方違反本協議之保密義務，應賠償他方因此所受之損害。', indent=0.7)
add_para('6.2　雙方同意，任一方如違反第二條或第三條之規定，應按次給付違約金新臺幣五十萬元；若違約情節連續或持續發生，以日計，每日另計新臺幣五萬元。前項違約金之給付，不影響他方依法律或本協議行使其他權利。', indent=0.7)
add_para('6.3　若實際損害逾前項違約金之金額，他方仍得請求超出部分。', indent=0.7)
add_para('6.4　損害賠償之範圍包含直接損失、調查及取證費用、以及合理之律師費與訴訟費用。', indent=0.7)

add_section_title('第七條　管轄與準據法')
add_para('7.1　本協議之解釋、效力及履行，以中華民國法律為準據法。', indent=0.7)
add_para('7.2　因本協議所生之爭議，雙方合意以臺灣臺北地方法院為第一審管轄法院。', indent=0.7)

add_section_title('第八條　其他條款')
add_para('8.1　本協議構成雙方就保密事項之完整合意，取代先前行為或口頭約定。', indent=0.7)
add_para('8.2　本協議之修改應經雙方書面同意。', indent=0.7)
add_para('8.3　本協議任一條款若被認定為無效或無法執行，不影響其他條款之效力。', indent=0.7)

add_para('')
add_section_title('雙方簽署欄')
add_para('甲方：（簽章）')
add_para('代表人：【                        】')
add_para('統一編號：【                        】')
add_para('地址：【                        】')
add_para('日期：中華民國　　　年　　　月　　　日')
add_para('')
add_para('乙方：（簽章）')
add_para('代表人：【                        】')
add_para('統一編號：【                        】')
add_para('地址：【                        】')
add_para('日期：中華民國　　　年　　　月　　　日')

# ── page number footer ──
footer = sec.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
from docx.oxml import OxmlElement
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), ' PAGE ')
run = OxmlElement('w:r')
t = OxmlElement('w:t')
t.text = '- 1 -'
run.append(t)
fld.append(run)
fp._p.append(fld)

OUTPUT = r'C:\Users\ysga1\Desktop\NDA_v1.2_FINAL.docx'
doc.save(OUTPUT)
print(f'✅ 產出：{OUTPUT}')
print(f'   格式：A4 / 邊界 2.5cm / 標楷體 14pt / 行高 28pt')
print(f'   頁數：{len(doc.paragraphs)} 段')
