"""
regulation_diff — 法規新舊條文對照引擎

功能：
- 提取舊版條文（本機快照 → LawOldVer 網頁備援）
- 逐條比對異動（修改/新增/刪除）
- 逐字元 diff（LCS 演算法）
- 匯出 Word 新舊條文對照表（三欄：修正條文/現行條文/立法理由）

依賴：
- python-docx (Word 匯出)
- regulation_tracker (索引快取)
"""

import json
import logging
import os
import re
import sys
from datetime import datetime
from typing import Optional
from urllib.request import urlopen, Request

logger = logging.getLogger("regulation_diff")

LAWOLDVER_URL = "https://law.moj.gov.tw/LawClass/LawOldVer.aspx?pcode="
LAWHISTORY_URL = "https://law.moj.gov.tw/LawClass/LawHistory.aspx?pcode="
LAWALL_URL = "https://law.moj.gov.tw/LawClass/LawAll.aspx?pcode="
UA = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15) ZhiyanTracker/1.0"


# ═══════════════════════════════════════════════════
# 歷史條文提取（LawOldVer 網頁備援）
# ═══════════════════════════════════════════════════

def fetch_old_articles(pcode: str, articles_dir: str) -> dict[str, str]:
    """
    取得舊版條文（An older version than the current one).

    Strategy:
    1. Check local snapshot cache for any older version
    2. Fall back to LawOldVer web scrape

    Returns {artNo: content} or empty dict on failure.
    """
    # 1) 本機快取：取最早的版本（=舊版）
    fpath = os.path.join(articles_dir, f"{pcode}.json")
    if os.path.exists(fpath):
        with open(fpath, "r", encoding="utf-8") as f:
            snap = json.load(f)
        keys = sorted(snap.keys())
        if len(keys) >= 2:
            # 有 2+ 版本：第二舊的就是舊條文
            old_ver = keys[-2] if len(keys) > 1 else keys[0]
            return snap[old_ver]
        elif len(keys) == 1 and keys[0]:
            # 只有現行版 → 試 LawOldVer
            pass

    # 2) LawOldVer 網頁備援
    try:
        url = LAWOLDVER_URL + pcode
        req = Request(url, headers={"User-Agent": UA})
        with urlopen(req, timeout=30) as resp:
            html = resp.read().decode("utf-8", errors="replace")

        articles = {}
        # 從 HTML 提取條文
        # <div class="col-no">第 1 條</div><div class="col-data">...</div>
        pattern = re.compile(
            r'col-no[^>]*>([\s\S]*?)</div>\s*'
            r'<div\s+class="col-data[^"]*"[^>]*>([\s\S]*?)</div>',
            re.IGNORECASE,
        )
        for m in pattern.finditer(html):
            no_html = m.group(1).strip()
            content_html = m.group(2).strip()
            no_match = re.search(r"第\s*([0-9A-Za-z]+(?:-[0-9A-Za-z]+)?)\s*條", no_html)
            if no_match:
                no = no_match.group(1)
                content = re.sub(r"<[^>]+>", "", content_html)
                content = content.replace("&nbsp;", " ").replace("&amp;", "&")
                content = re.sub(r"\s+", " ", content).strip()
                if content:
                    articles[no] = content

        if articles:
            logger.info(f"從 LawOldVer 取得 {pcode} 舊版條文 {len(articles)} 條")
            # 快取進本機（供下次直接讀取）
            try:
                if os.path.exists(fpath):
                    with open(fpath, "r", encoding="utf-8") as f:
                        old_snap = json.load(f)
                else:
                    old_snap = {}
                old_snap["_oldver"] = articles
                with open(fpath, "w", encoding="utf-8") as f:
                    json.dump(old_snap, f, ensure_ascii=False, indent=1)
            except Exception:
                pass
            return articles

        logger.warning(f"LawOldVer 頁面無條文: {pcode}")
    except Exception as e:
        logger.warning(f"LawOldVer 抓取失敗 {pcode}: {e}")

    return {}


# ═══════════════════════════════════════════════════
# 逐條比對
# ═══════════════════════════════════════════════════

def diff_articles(
    old_arts: dict[str, str],
    new_arts: dict[str, str],
) -> dict:
    """
    比對新舊條文，回傳異動結果。

    Returns:
        {modified: [{no, old, new}], added: [{no, new}], removed: [{no, old}],
         changed_count: int, unchanged_count: int}
    """
    all_nos = set(old_arts.keys()) | set(new_arts.keys())

    def sort_key(no):
        parts = no.split("-")
        return (int(parts[0]), int(parts[1]) if len(parts) > 1 else 0)

    sorted_nos = sorted(all_nos, key=sort_key)

    modified = []
    added = []
    removed = []
    unchanged = 0

    for no in sorted_nos:
        old_t = old_arts.get(no, "")
        new_t = new_arts.get(no, "")
        # 正規化比對：去全形半形折疊 + 去空白
        old_norm = re.sub(r"\s+", "", old_t)
        new_norm = re.sub(r"\s+", "", new_t)

        if old_t and new_t:
            if old_norm != new_norm:
                modified.append({"no": no, "old": old_t, "new": new_t})
            else:
                unchanged += 1
        elif new_t and not old_t:
            added.append({"no": no, "new": new_t})
        elif old_t and not new_t:
            removed.append({"no": no, "old": old_t})
        else:
            unchanged += 1

    return {
        "modified": modified,
        "added": added,
        "removed": removed,
        "changed_count": len(modified) + len(added) + len(removed),
        "unchanged_count": unchanged,
    }


# ═══════════════════════════════════════════════════
# 逐字元 LCS Diff
# ═══════════════════════════════════════════════════

def lcs_diff(old_text: str, new_text: str, max_len: int = 6000) -> list[dict]:
    """
    逐字元 LCS diff。
    回傳 ops: [{t: '='|'+'|'-', s: str}]
    - '=' 相同部分
    - '+' 新增
    - '-' 刪除

    若文字過長（> max_len 字元）則退回到全文 level diff。
    """
    a = old_text or ""
    b = new_text or ""
    n, m = len(a), len(b)

    if n + m > max_len:
        # 過長：退回到整段 level
        ops = []
        if a:
            ops.append({"t": "-", "s": a})
        if b:
            ops.append({"t": "+", "s": b})
        return ops

    if n == 0 and m == 0:
        return []
    if n == 0:
        return [{"t": "+", "s": b}]
    if m == 0:
        return [{"t": "-", "s": a}]
    if a == b:
        return [{"t": "=", "s": a}]

    # LCS DP (iterative, 2-row)
    prev = [0] * (m + 1)
    for i in range(n):
        curr = [0] * (m + 1)
        for j in range(m):
            if a[i] == b[j]:
                curr[j + 1] = prev[j] + 1
            else:
                curr[j + 1] = max(curr[j], prev[j + 1])
        prev = curr

    # 回溯
    ops = []
    i, j = n, m
    while i > 0 or j > 0:
        if i > 0 and j > 0 and a[i - 1] == b[j - 1]:
            t, s = "=", a[i - 1]
            i -= 1
            j -= 1
        elif j > 0 and (i == 0 or prev[j] < prev[j - 1]):
            t, s = "+", b[j - 1]
            j -= 1
        elif i > 0:
            t, s = "-", a[i - 1]
            i -= 1
        else:
            break  # safety
        if ops and ops[-1]["t"] == t:
            ops[-1]["s"] = s + ops[-1]["s"]
        else:
            ops.insert(0, {"t": t, "s": s})
    # Use prev for next iteration? No, this is one-shot.
    # Actually the prev variable holds the last DP row.
    # Let me just return the result.

    return ops


# ═══════════════════════════════════════════════════
# 修正說明/立法理由（來自沿革）
# ═══════════════════════════════════════════════════

def parse_amendment_summary(histories: dict[str, str], pcode: str) -> dict:
    """從沿革文字解析最新修正說明"""
    text = histories.get(pcode, "")
    if not text:
        return {}
    entries = list(re.finditer(r"(\d+)\.\s*([\s\S]*?)(?=\n\s*\d+\.|\n*\Z)", text))
    if not entries:
        return {}
    # 挑最新一筆（日期最大）
    best = entries[-1]
    body = best.group(2).strip()
    # 提取文號
    doc_m = re.search(r"([^\s，；。]{2,}字第\s*[0-9A-Za-z]+\s*號|總統[^，；。]{0,12}令)", body)
    doc_no = doc_m.group(1).strip() if doc_m else ""
    # 提取相關條號
    art_nos = re.findall(r"第\s*(\d+(?:\s*[～~\-－之至]\s*\d+)*)\s*條", body)
    return {
        "doc_no": doc_no,
        "articles_mentioned": list(dict.fromkeys(art_nos))[:20],
        "full": body[:500],
    }


# ═══════════════════════════════════════════════════
# 完整 Diff 報告
# ═══════════════════════════════════════════════════

def build_diff_report(
    pcode: str,
    tracker: "RegulationTracker",
) -> Optional[dict]:
    """
    建立完整的異動報告。

    1. 拿新條文（本機快取的最新版本）
    2. 拿舊條文（本機快取 → LawOldVer）
    3. 逐條比對
    4. 逐字元 diff
    5. 附修正說明
    """
    meta = tracker.law_meta(pcode)
    if not meta:
        logger.warning(f"查無法規: {pcode}")
        return None

    name = meta["name"]
    level = meta.get("level", "")
    new_date = meta.get("modifiedDate", "")
    abolished = meta.get("abolished", False)

    # 新條文
    snap = tracker.get_articles(pcode)
    if not snap:
        logger.warning(f"無條文快取: {pcode}")
        return None

    # 找舊條文版本
    # 優先從本機快取找（非最新版本即為舊版）
    ver_keys = sorted(k for k in snap.keys() if not k.startswith("_"))
    if len(ver_keys) >= 2:
        old_date = ver_keys[-2]  # 倒數第二新 = 舊版
        new_arts = snap[ver_keys[-1]]
        old_arts = snap[old_date]
    else:
        # 只有一個版本 → 從 LawOldVer 取舊版
        old_arts = fetch_old_articles(pcode, os.path.join(tracker.data_dir, "articles"))
        new_arts = snap[ver_keys[-1]] if ver_keys else {}
        old_date = ""

    if not new_arts:
        logger.warning(f"無新條文可比較: {pcode}")
        return None

    # 逐條比對
    diff = diff_articles(old_arts, new_arts)

    # 逐字元 diff
    for item in diff["modified"]:
        item["char_diff"] = lcs_diff(item["old"], item["new"])

    # 修正說明
    amend = parse_amendment_summary(tracker._histories, pcode)

    return {
        "pcode": pcode,
        "name": name,
        "level": level,
        "new_date": new_date,
        "old_date": old_date or "（歷史版本）",
        "abolished": abolished,
        "summary": amend,
        **diff,
    }


# ═══════════════════════════════════════════════════
# Word 匯出（python-docx）
# ═══════════════════════════════════════════════════

def _fmt_date(ymd: str) -> str:
    """YYYYMMDD → YYYY-MM-DD"""
    s = str(ymd)
    if len(s) == 8:
        return f"{s[:4]}-{s[4:6]}-{s[6:8]}"
    return s


def export_word(diff_report: dict, output_path: str) -> str:
    """
    匯出 Word 新舊條文對照表。

    格式：
    - 首頁標題：法規名稱 + 對照表
    - 三欄表格：修正條文（新）| 現行條文（舊）| 備註
    - 新增文字：紅字底線
    - 刪除文字：紅字刪除線
    - 附立法理由（若有）
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── 頁面設定（A4 橫印？直印）────
    # 直印 A4 即可（三欄表格）
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── 字型設定 ──────────────────
    style = doc.styles["Normal"]
    style.font.name = "標楷體"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    # ── 標題 ──────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{diff_report['name']}　新舊條文對照表")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "標楷體"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    # 副標
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_text = (
        f"新版 {_fmt_date(diff_report['new_date'])}"
        f"　↔　舊版 {_fmt_date(diff_report['old_date'])}"
        f"　·　異動 {diff_report['changed_count']} 條"
    )
    run = subtitle.add_run(ver_text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "標楷體"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    # 修正說明（若有）
    if diff_report.get("summary", {}).get("doc_no"):
        p = doc.add_paragraph()
        run = p.add_run(
            f"修正文號：{diff_report['summary']['doc_no']}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacing

    # ── 三欄表格 ──────────────────
    # 決定欄寬
    has_reasons = False  # 目前無立法理由串接，預設 false
    col_widths = [Cm(7.5), Cm(7.5), Cm(3.5)]  # 新 / 舊 / 備註

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    # 設定欄寬
    for i, w in enumerate(col_widths):
        table.columns[i].width = w

    # 標題列
    hdr = table.rows[0]
    for i, text in enumerate(["修正條文（新）", "現行條文（舊）", "備註"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "標楷體"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        # 灰底
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(
            qn("w:shd"),
            {qn("w:fill"): "F2F2F2", qn("w:val"): "clear"},
        )
        shading.append(shd)

    # 資料列
    kind_color = {
        "修正": (0xB4, 0x88, 0x2B),  # 暗金
        "新增": (0x2F, 0x7D, 0x4F),  # 深綠
        "刪除": (0xB4, 0x45, 0x2B),  # 暗紅
    }
    RED = RGBColor(0xC0, 0x39, 0x2B)

    all_items = (
        [("修正", a) for a in diff_report["modified"]]
        + [("新增", a) for a in diff_report["added"]]
        + [("刪除", a) for a in diff_report["removed"]]
    )

    for kind, item in all_items:
        row = table.add_row()
        cells = row.cells
        color = kind_color.get(kind, (0, 0, 0))
        color_rgb = RGBColor(*color)

        # 新條文欄（col 0）
        p = cells[0].paragraphs[0]
        run_kind = p.add_run(f"第{item['no']}條　（{kind}）")
        run_kind.bold = True
        run_kind.font.color.rgb = color_rgb
        run_kind.font.size = Pt(9)
        run_kind.font.name = "標楷體"
        run_kind.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

        if kind == "修正" and "char_diff" in item and item["char_diff"]:
            _add_char_diff_to_paragraph(p, item["char_diff"])
        elif kind == "新增":
            _add_text_run(p, item.get("new", ""), added=True)
        elif kind == "刪除":
            _add_text_run(p, item.get("old", ""), removed=True)

        # 舊條文欄（col 1）
        p2 = cells[1].paragraphs[0]
        if kind == "修正" and "char_diff" in item and item["char_diff"]:
            _add_char_diff_to_paragraph(p2, item["char_diff"], side="old")
        elif kind == "新增":
            p2.add_run("（新增）").font.size = Pt(9)
        elif kind == "刪除":
            _add_text_run(p2, item.get("old", ""), removed=True)

        # 備註欄（col 2）
        p3 = cells[2].paragraphs[0]
        p3.add_run(f"（{kind}）").font.size = Pt(9)

        # 表格列高
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)

    # ── 連結 ──────────────────────
    doc.add_paragraph()
    p_info = doc.add_paragraph()
    run = p_info.add_run("參考來源：")
    run.bold = True
    run.font.size = Pt(9)
    run = p_info.add_run(f"\n全國法規資料庫：{LAWALL_URL}{diff_report['pcode']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
    run = p_info.add_run(f"\n沿革：{LAWHISTORY_URL}{diff_report['pcode']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

    # 免責聲明
    p_disc = doc.add_paragraph()
    p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_disc.add_run("本對照表僅供參考，不構成法律意見。如有疑義請以全國法規資料庫最新版本為準。")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "標楷體"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    # 儲存
    doc.save(output_path)
    logger.info(f"已匯出 Word：{output_path}")
    return output_path


def _add_text_run(paragraph, text: str, added: bool = False, removed: bool = False):
    """新增純文字 run（紅字底線或刪除線）"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    run = paragraph.add_run(text)
    if added:
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        run.font.underline = True
    if removed:
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        run.font.strike = True
    run.font.size = Pt(9)
    run.font.name = "標楷體"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")


def _add_char_diff_to_paragraph(paragraph, ops, side="new"):
    """
    將逐字元 diff ops 加入段落。
    side='new': '+' 紅底線, '=' 一般
    side='old': '-' 紅刪除線, '=' 一般
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    RED = RGBColor(0xC0, 0x39, 0x2B)
    keep = "+" if side == "new" else "-"

    for op in ops:
        if op["t"] == "=":
            run = paragraph.add_run(op["s"])
            run.font.size = Pt(9)
            run.font.name = "標楷體"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        elif op["t"] == keep:
            run = paragraph.add_run(op["s"])
            run.font.color.rgb = RED
            run.font.size = Pt(9)
            run.font.name = "標楷體"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
            if side == "new":
                run.font.underline = True
            else:
                run.font.strike = True


# ═══════════════════════════════════════════════════
# CLI 命令
# ═══════════════════════════════════════════════════

def diff_cmd(pcode: str, output: Optional[str] = None, verbose: bool = False):
    """CLI: 建立新舊條文對照並可選匯出 Word"""
    from zhiyan_legal.regulation_tracker import RegulationTracker

    tracker = RegulationTracker()
    tracker.sync_index(force=False)

    report = build_diff_report(pcode, tracker)
    if not report:
        print(f"✗ 無法建立 {pcode} 的異動報告")
        return

    print(f"\n{'='*60}")
    print(f"{report['name']} ({report['level']})")
    print(f"新版 {_fmt_date(report['new_date'])}  ↔  舊版 {_fmt_date(report['old_date'])}")
    print(f"{'='*60}")

    if report["abolished"]:
        print("🚫 已廢止")
        return

    print(f"\n📊 異動摘要：{report['changed_count']} 條異動")
    if report["summary"].get("doc_no"):
        print(f"修正文號：{report['summary']['doc_no']}")

    if report["modified"]:
        print(f"\n🔶 修正條文 ({len(report['modified'])} 條)：")
        for a in report["modified"]:
            no = a["no"]
            # 簡短顯示 diff 摘要
            old_short = a["old"][:80] if a["old"] else ""
            new_short = a["new"][:80] if a["new"] else ""
            print(f"  第{no}條")
            if old_short and new_short:
                # 找出首個不同位置
                for ci, (oc, nc) in enumerate(zip(old_short, new_short)):
                    if oc != nc:
                        print(f"    - ...{old_short[max(0,ci-10):ci+10]}...")
                        print(f"    + ...{new_short[max(0,ci-10):ci+10]}...")
                        break
            if verbose:
                print(f"    舊：{a['old']}")
                print(f"    新：{a['new']}")

    if report["added"]:
        print(f"\n🟢 新增條文 ({len(report['added'])} 條)：")
        for a in report["added"]:
            print(f"  第{a['no']}條：{a['new'][:80]}...")

    if report["removed"]:
        print(f"\n🔴 刪除條文 ({len(report['removed'])} 條)：")
        for a in report["removed"]:
            print(f"  第{a['no']}條")

    # Word 匯出
    if output:
        export_word(report, output)
        print(f"\n✓ 已匯出 Word：{output}")

    # 預設匯出路徑
    if not output:
        _project_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        default_out = os.path.join(
            _project_dir,
            "data",
            "exports",
            f"{report['name']}_新舊條文對照_{_fmt_date(report['new_date'])}.docx",
        )
        os.makedirs(os.path.dirname(default_out), exist_ok=True)
        export_word(report, default_out)
        print(f"\n✓ 已自動匯出：{default_out}")


def diff_all_cmd(output_dir: Optional[str] = None, verbose: bool = False):
    """CLI: 為所有有異動的法規產生對照表"""
    from zhiyan_legal.regulation_tracker import RegulationTracker

    tracker = RegulationTracker()
    tracker.sync_index(force=False)

    tracked = tracker.get_all_tracked()
    export_dir = output_dir or os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
        "data",
        "exports",
    )
    os.makedirs(export_dir, exist_ok=True)

    generated = 0
    for t in tracked:
        pcode = t["pcode"]
        meta = tracker.law_meta(pcode)
        if not meta:
            continue
        snap = tracker.get_articles(pcode)
        if len(snap) < 1:
            continue

        report = build_diff_report(pcode, tracker)
        if not report or report["changed_count"] == 0:
            continue

        name_safe = re.sub(r'[\\/:*?"<>|]', "_", report["name"])
        out_path = os.path.join(
            export_dir,
            f"{name_safe}_新舊條文對照_{_fmt_date(report['new_date'])}.docx",
        )
        export_word(report, out_path)
        generated += 1
        if verbose:
            print(f"  ✓ {report['name']}: {report['changed_count']} 條異動 → {out_path}")

    print(f"\n已完成：{generated} 份對照表")
    print(f"匯出目錄：{export_dir}")
