"""
Zhiyan Legal — Court-compliant document generator.

Generates .docx files that conform to:
- 民事訴訟書狀規則 §3 (2025/08/29 amendment)
- 憲法法庭書狀規則
- 行政訴訟書狀規則

Usage:
    python -c "
    from zhiyan_legal.doc_generator import LegalDocument
    doc = LegalDocument()
    doc.add_title('刑事聲請延緩執行狀')
    doc.add_reference('案號：XXX年度XXX字第XXX號')
    doc.add_body('聲請事項：\n一、...\n二、...')
    doc.add_body('事實及理由：\n...')
    doc.save('/tmp/聲請狀.docx')
    "
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── 民事訴訟書狀規則 §3 強制參數 ─────────────────────────

A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
MARGIN_CM = 2.5          # 上下左右 2.5cm
FONT_NAME = "TW-Kai"      # 標楷體（全字庫正楷體）
FONT_SIZE_PT = 14         # 14pt
LINE_HEIGHT_PT = 28       # 行高 28pt（25~30pt 範圍內，閱讀舒適）
MAX_FONT_SIZE_PT = 20     # 上限 20pt
MIN_FONT_SIZE_PT = 14     # 下限 14pt


class LegalDocument:
    """法院合規書狀產生器。"""
    
    def __init__(self, filename: str = ""):
        self.doc = Document()
        self._setup_page()
        self._setup_default_style()
    
    def _setup_page(self):
        """設定 A4 頁面 + 2.5cm 邊界。"""
        section = self.doc.sections[0]
        section.page_width = Cm(A4_WIDTH_CM)
        section.page_height = Cm(A4_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
    
    def _setup_default_style(self):
        """設定預設段落樣式：標楷體 14pt，固定行高 28pt。"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = Pt(FONT_SIZE_PT)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        
        pf = style.paragraph_format
        pf.line_spacing = Pt(LINE_HEIGHT_PT)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
    
    def add_title(self, text: str, font_size: int = 18):
        """加入書狀標題（置中，18pt 或自訂，上限 20pt）。"""
        fs = min(font_size, MAX_FONT_SIZE_PT)
        fs = max(fs, MIN_FONT_SIZE_PT)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(fs)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        return p
    
    def add_subtitle(self, text: str):
        """加入副標題（置中，14pt）。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        return p
    
    def add_reference(self, text: str):
        """加入案號/受文者（靠左，14pt）。"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        return p
    
    def add_section_title(self, text: str, font_size: int = 16):
        """加入段落標題（如 聲請事項、事實及理由）。"""
        fs = min(font_size, MAX_FONT_SIZE_PT)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(fs)
        run.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        return p
    
    def add_body(self, text: str):
        """加入內文（14pt，固定行高 28pt，雙行標楷體）。"""
        paragraphs = text.strip().split('\n')
        last_p = None
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            p = self.doc.add_paragraph()
            # 首行縮排 2 字元
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0.7)
            run = p.add_run(para_text.strip())
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_PT)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            last_p = p
        return last_p
    
    def add_indent_body(self, text: str, level: int = 1):
        """加入縮排內文（層級用於條列項目）。"""
        indent_cm = 0.7 + (level - 1) * 0.7
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(indent_cm)
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        return p
    
    def add_page_number(self):
        """在頁面底端置中加入頁碼（符合 §3 第三款）。"""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = p.add_run()
        fld_xml = (
            '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:instr=" PAGE "><w:r><w:t>- 1 -</w:t></w:r></w:fldSimple>'
        )
        from lxml import etree
        run._r.append(etree.fromstring(fld_xml))
    
    def save(self, path: str):
        """輸出 .docx 檔案。"""
        self.add_page_number()
        self.doc.save(path)
        print(f"✅ 書狀已產生：{path}")
        print(f"   格式：A4 / 邊界 2.5cm / 標楷體 {FONT_SIZE_PT}pt / 行高 {LINE_HEIGHT_PT}pt")
